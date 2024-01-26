from docx import Document

# Create a new Word document
doc = Document()

# Define the content based on the user's requirements
content = [
    ("1. Redefining Work and Success", 
    "Concept: Shifting from a traditional success model focused on high income and long hours to one that values relative income, personal freedom, and overall quality of life.",
    [
        ("1. Personal Vision and Goal Setting:", 
        ["A. Self-Reflection: Spend time reflecting on what brings you true happiness and fulfillment. This could involve journaling, meditation, or even professional coaching.",
         "B. Goal Frameworks: Use frameworks like SMART (Specific, Measurable, Achievable, Relevant, Time-bound) goals to set clear and actionable objectives in various life domains.",
         "C. Vision Boards: Create a vision board that visually represents your goals and aspirations to keep you motivated and focused."]),
        ("2. Lifestyle-Centric Financial Planning:", 
        ["A. Expense Analysis: Regularly review and categorize your expenses to identify areas where you can cut back to fund your desired lifestyle.",
         "B. Passive Income Strategies: Research and invest in passive income streams. This could include rental properties, dividend-yielding stocks, or creating digital products.",
         "C. Financial Education: Continuously educate yourself on financial management and investment strategies to make informed decisions."]),
        ("3. Mindset and Cultural Shift:", 
        ["A. Community Engagement: Join online forums, local clubs, or groups that share your new perspective on work and success.",
         "B. Mindset Training: Practice mindfulness or other mental training techniques to reinforce your new mindset.",
         "C. Inspirational Content: Regularly consume books, podcasts, and other media that reinforce your values and goals."])
    ]),
    ("2. Efficiency and Effectiveness", 
    "Concept: Prioritizing the most impactful activities for effectiveness and streamlining processes for efficiency.",
    [
        ("1. Regular Task Audits:", 
        ["A. Task Inventory: Make a comprehensive list of all your regular tasks and responsibilities.",
         "B. Impact Analysis: Evaluate each task based on its impact towards your goals and its alignment with your strengths.",
         "C. Delegation Plan: Identify tasks that can be delegated and develop a plan for training others to take them over."]),
        ("2. Adopting Productivity Techniques:", 
        ["A. Eisenhower Box: Categorize tasks into urgent/non-urgent and important/non-important to prioritize effectively.",
         "B. Time Blocking: Allocate specific blocks of time for focused work, and use tools like digital calendars for scheduling.",
         "C. Productivity Apps: Utilize apps like Todoist or Trello to manage and track your tasks efficiently."]),
        ("3. Leveraging Technology:", 
        ["A. Automation Tools: Explore tools like Zapier or IFTTT for automating repetitive digital tasks.",
         "B. Project Management Software: Use software like Asana or Monday.com to streamline workflow and collaboration.",
         "C. AI and Machine Learning: Investigate AI tools that can automate complex tasks like data analysis or customer service."])
    ]),
    ("3. Automation and Outsourcing", 
    "Concept: Using technology and external resources to handle routine tasks, freeing up time for more meaningful work.",
    [
        ("1. Identify Automation Opportunities:", 
        ["A. Task Mapping: Map out your daily routines and identify repetitive or time-consuming tasks.",
         "B. Automation Research: Stay informed about the latest automation technologies relevant to your field or daily life.",
         "C. Trial and Error: Experiment with different automation tools to find what works best for you."]),
        ("2. Smart Outsourcing:", 
        ["A. Task Assessment: Determine which tasks require specialized skills or can be done more cost-effectively by others.",
         "B. Vendor Selection: Research and select outsourcing partners or platforms based on reliability, cost, and quality of work.",
         "C. Communication and Management: Establish clear communication channels and expectations with your outsourcing partners."]),
        ("3. Process Optimization:", 
        ["A. Standard Operating Procedures (SOPs): Develop SOPs for tasks you outsource to ensure consistency and ease of training.",
         "B. Feedback Loops: Implement regular feedback mechanisms to continuously improve the processes.",
         "C. Performance Metrics: Set and monitor performance metrics for outsourced tasks to ensure they meet your standards."])
    ])
]

# Add content to the document
for section_title, section_concept, section_points in content:
    doc.add_heading(section_title, level=2)
    doc.add_paragraph(section_concept)

    for point_title, subpoints in section_points:
        # Add main points (formatted bold)
        p = doc.add_paragraph()
        p.add_run(point_title).bold = True

        # Add subpoints and bold the initial part before the colon
        for subpoint in subpoints:
            p = doc.add_paragraph()
            bold_part, rest_of_text = subpoint.split(':', 1)
            p.add_run(bold_part + ':').bold = True
            p.add_run(rest_of_text)
# Save the document
doc.save('your_document.docx')